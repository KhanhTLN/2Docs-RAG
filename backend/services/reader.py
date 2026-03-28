from __future__ import annotations
import re, unicodedata, logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)


@dataclass
class Paragraph:
    text:  str
    style: str = "Normal"   # "Heading 1" / "Heading 2" / "Normal"
    page:  int = 0


@dataclass
class ParsedDoc:
    name:       str
    fmt:        str                           # "docx" | "pdf"
    paragraphs: List[Paragraph] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs)


class DocumentReader:
    """Đọc 1 file DOCX hoặc PDF → ParsedDoc."""

    def read(self, file_path: str) -> ParsedDoc:
        path = Path(file_path)
        ext  = path.suffix.lower()
        if ext == ".docx":
            return self._docx(path)
        if ext == ".pdf":
            return self._pdf(path)
        raise ValueError(f"Không hỗ trợ định dạng: {ext}")

    # ── DOCX ──────────────────────────────────────────────────────────

    def _docx(self, path: Path) -> ParsedDoc:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("pip install python-docx")

        doc   = Document(str(path))
        norm  = _Normalizer()
        paras = []
        for p in doc.paragraphs:
            text = norm.clean(p.text)
            if len(text) >= 2:
                paras.append(Paragraph(text=text, style=p.style.name, page=0))

        logger.info(f"DOCX: {path.name} → {len(paras)} đoạn")
        return ParsedDoc(name=path.name, fmt="docx", paragraphs=paras)

    # ── PDF ───────────────────────────────────────────────────────────

    def _pdf(self, path: Path) -> ParsedDoc:
        import fitz
        norm, paras = _Normalizer(), []
        doc = fitz.open(str(path))
        for page_num, page in enumerate(doc, 1):
            for block in page.get_text("blocks", sort=True):
                if block[6] != 0: continue
                for line in block[4].split("\n"):
                    text = norm.clean(line)
                    if len(text) > 5:
                        paras.append(Paragraph(text=text, style=_guess_style(text), page=page_num))
        doc.close()
        return ParsedDoc(name=path.name, fmt="pdf", paragraphs=paras)


class _Normalizer:
    _OCR = {
        r"\bĐIỀU\b": "Điều", r"\bĐiêu\b": "Điều", r"\bĐieu\b": "Điều",
        r"\bKHOẢN\b": "Khoản", r"\bKhoan\b": "Khoản",
        r"\bĐIỂM\b": "Điểm", r"\bPHỤ LỤC\b": "Phụ lục",
    }

    def clean(self, text: str) -> str:
        text = unicodedata.normalize("NFC", text)
        text = re.sub(r"-\s*\d+\s*-", "", text)             # số trang
        text = re.sub(r"[Tt]rang\s+\d+\s*/\s*\d+", "", text)
        for pat, rep in self._OCR.items():
            text = re.sub(pat, rep, text)
        text = re.sub(r"\t+", " ", text)
        text = re.sub(r" {2,}", " ", text)
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = text.replace("\u2018", "'").replace("\u2019", "'")
        text = text.replace("\u2013", "-").replace("\u2014", "-")
        return text.strip()


def _guess_style(text: str) -> str:
    if re.match(r"^(Điều|ĐIỀU)\s+\d+", text):
        return "Heading 1"
    if re.match(r"^\d+[\.\)]\s+\S", text):
        return "Heading 2"
    return "Normal"
